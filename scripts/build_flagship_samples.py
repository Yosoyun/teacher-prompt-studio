#!/usr/bin/env python3
"""Build the verified bilingual Class 10 flagship assessment sources.

Final public files are produced from these DOCX sources with LibreOffice and
rendered with Poppler during release QA. This script intentionally keeps
student and teacher material in physically separate documents.
"""

from __future__ import annotations

from collections import Counter
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
WORK_DIR = ROOT / "tmp" / "sample_artifacts"
PUBLIC_DIR = ROOT / "public" / "samples"
STUDENT_DOCX = PUBLIC_DIR / "class-10-quadratics-student-editable.docx"
TEACHER_SOURCE = WORK_DIR / "class-10-quadratics-teacher-pack-source.docx"

LATIN_FONT = "Arial"
DEVANAGARI_FONT = "Noto Sans Devanagari"
INK = "102A25"
INK_SOFT = "36514B"
BLUE = "3457E8"
LIME = "D9FF64"
MINT = "EAF8F4"
PALE_BLUE = "EEF2FF"
PALE_GOLD = "FFF6DE"
WHITE = "FFFFFF"
LINE = "CFD8D4"
MUTED = "66756F"


ITEMS = [
    {
        "n": 1, "section": "mcq", "marks": 1,
        "en": "The roots of x² - 5x + 6 = 0 are:",
        "hi": "x² - 5x + 6 = 0 के मूल हैं:",
        "options": ["A. 1 and 6", "B. 2 and 3", "C. -2 and -3", "D. -1 and -6"],
        "answer": "B",
        "solution_en": "x² - 5x + 6 = (x - 2)(x - 3), so the roots are 2 and 3.",
        "solution_hi": "x² - 5x + 6 = (x - 2)(x - 3), इसलिए मूल 2 और 3 हैं।",
        "marking": "1 mark for option B.",
    },
    {
        "n": 2, "section": "mcq", "marks": 1,
        "en": "The discriminant of x² + 4x + 5 = 0 is:",
        "hi": "x² + 4x + 5 = 0 का विविक्तकर (discriminant) है:",
        "options": ["A. 36", "B. 16", "C. 0", "D. -4"],
        "answer": "D",
        "solution_en": "D = b² - 4ac = 4² - 4(1)(5) = -4.",
        "solution_hi": "D = b² - 4ac = 4² - 4(1)(5) = -4।",
        "marking": "1 mark for option D.",
    },
    {
        "n": 3, "section": "mcq", "marks": 1,
        "en": "The sum of the roots of 2x² - 7x + 3 = 0 is:",
        "hi": "2x² - 7x + 3 = 0 के मूलों का योग है:",
        "options": ["A. 7", "B. -7/2", "C. 7/2", "D. 3/2"],
        "answer": "C",
        "solution_en": "For ax² + bx + c = 0, sum of roots = -b/a = -(-7)/2 = 7/2.",
        "solution_hi": "ax² + bx + c = 0 में मूलों का योग -b/a है; अतः -(-7)/2 = 7/2।",
        "marking": "1 mark for option C.",
    },
    {
        "n": 4, "section": "mcq", "marks": 1,
        "en": "Which equation has -3 and 4 as its roots?",
        "hi": "निम्न में से किस समीकरण के मूल -3 और 4 हैं?",
        "options": ["A. x² - x - 12 = 0", "B. x² + x - 12 = 0", "C. x² - 7x + 12 = 0", "D. x² + 7x + 12 = 0"],
        "answer": "A",
        "solution_en": "Sum = 1 and product = -12, so x² - (1)x - 12 = 0.",
        "solution_hi": "मूलों का योग 1 और गुणनफल -12 है; इसलिए समीकरण x² - x - 12 = 0 है।",
        "marking": "1 mark for option A.",
    },
    {
        "n": 5, "section": "mcq", "marks": 1,
        "en": "The vertex of the parabola y = (x - 2)² - 5 is:",
        "hi": "परवलय y = (x - 2)² - 5 का शीर्ष है:",
        "options": ["A. (-2, -5)", "B. (2, 5)", "C. (2, -5)", "D. (-5, 2)"],
        "answer": "C",
        "solution_en": "The equation is in vertex form y = (x - h)² + k, so (h, k) = (2, -5).",
        "solution_hi": "यह y = (x - h)² + k के रूप में है; अतः शीर्ष (h, k) = (2, -5) है।",
        "marking": "1 mark for option C.",
    },
    {
        "n": 6, "section": "mcq", "marks": 1,
        "en": "If the discriminant of a quadratic equation is 0, its roots are:",
        "hi": "यदि किसी द्विघात समीकरण का विविक्तकर 0 है, तो उसके मूल:",
        "options": ["A. real and equal", "B. real and unequal", "C. non-real", "D. always zero"],
        "answer": "A",
        "solution_en": "D = 0 gives one repeated real root, -b/(2a).",
        "solution_hi": "D = 0 होने पर -b/(2a) एक पुनरावृत्त वास्तविक मूल होता है।",
        "marking": "1 mark for option A.",
    },
    {
        "n": 7, "section": "mcq", "marks": 1,
        "en": "Which parabola opens downward?",
        "hi": "कौन-सा परवलय नीचे की ओर खुलता है?",
        "options": ["A. y = 2x² + 3x - 1", "B. y = x² - 8", "C. y = -2x² + x + 4", "D. y = 0.5x² - 3x"],
        "answer": "C",
        "solution_en": "A parabola opens downward when the coefficient of x² is negative.",
        "solution_hi": "x² का गुणांक ऋणात्मक होने पर परवलय नीचे की ओर खुलता है।",
        "marking": "1 mark for option C.",
    },
    {
        "n": 8, "section": "mcq", "marks": 1,
        "en": "A rectangle has width x cm, length (x + 2) cm and area 48 cm². Which equation models the situation?",
        "hi": "एक आयत की चौड़ाई x सेमी, लंबाई (x + 2) सेमी और क्षेत्रफल 48 सेमी² है। सही समीकरण कौन-सा है?",
        "options": ["A. x² + 2x - 48 = 0", "B. x² - 2x - 48 = 0", "C. 2x² + 2x - 48 = 0", "D. x² + 48x + 2 = 0"],
        "answer": "A",
        "solution_en": "Area = x(x + 2) = 48, hence x² + 2x - 48 = 0.",
        "solution_hi": "क्षेत्रफल x(x + 2) = 48 है; अतः x² + 2x - 48 = 0।",
        "marking": "1 mark for option A.",
    },
    {
        "n": 9, "section": "short", "marks": 2,
        "en": "Solve x² + x - 12 = 0 by factorisation.",
        "hi": "x² + x - 12 = 0 को गुणनखंड विधि से हल कीजिए।",
        "answer": "x = 3 or x = -4",
        "solution_en": "x² + x - 12 = (x + 4)(x - 3) = 0. Therefore x = -4 or x = 3.",
        "solution_hi": "x² + x - 12 = (x + 4)(x - 3) = 0। इसलिए x = -4 अथवा x = 3।",
        "marking": "1 mark for correct factorisation; 1 mark for both roots.",
    },
    {
        "n": 10, "section": "short", "marks": 2,
        "en": "Solve 2x² + x - 3 = 0 using the quadratic formula.",
        "hi": "2x² + x - 3 = 0 को द्विघात सूत्र से हल कीजिए।",
        "answer": "x = 1 or x = -3/2",
        "solution_en": "x = [-1 ± √(1 + 24)]/4 = (-1 ± 5)/4, giving x = 1 or x = -3/2.",
        "solution_hi": "x = [-1 ± √(1 + 24)]/4 = (-1 ± 5)/4; अतः x = 1 अथवा x = -3/2।",
        "marking": "1 mark for correct substitution; 1 mark for both roots.",
    },
    {
        "n": 11, "section": "short", "marks": 2,
        "en": "Find all values of k for which x² + kx + 16 = 0 has equal real roots.",
        "hi": "k के सभी मान ज्ञात कीजिए जिनके लिए x² + kx + 16 = 0 के बराबर वास्तविक मूल हों।",
        "answer": "k = 8 or k = -8",
        "solution_en": "Equal roots require D = 0: k² - 64 = 0, so k = ±8.",
        "solution_hi": "बराबर मूलों के लिए D = 0: k² - 64 = 0; इसलिए k = ±8।",
        "marking": "1 mark for k² - 64 = 0; 1 mark for k = ±8.",
    },
    {
        "n": 12, "section": "short", "marks": 2,
        "en": "Form a monic quadratic polynomial whose roots are 2 + √3 and 2 - √3.",
        "hi": "2 + √3 और 2 - √3 मूलों वाला एक एकक-अग्र-गुणांक द्विघात बहुपद बनाइए।",
        "answer": "x² - 4x + 1",
        "solution_en": "Sum = 4 and product = (2 + √3)(2 - √3) = 1. Polynomial: x² - 4x + 1.",
        "solution_hi": "मूलों का योग 4 और गुणनफल 1 है। अतः बहुपद x² - 4x + 1 है।",
        "marking": "1 mark for sum/product; 1 mark for the polynomial.",
    },
    {
        "n": 13, "section": "short", "marks": 2,
        "en": "Use the discriminant to explain why x² - 2x + 5 = 0 has no real roots.",
        "hi": "विविक्तकर का उपयोग करके समझाइए कि x² - 2x + 5 = 0 के कोई वास्तविक मूल क्यों नहीं हैं।",
        "answer": "D = -16; therefore there are no real roots",
        "solution_en": "D = (-2)² - 4(1)(5) = 4 - 20 = -16. Since D < 0, the roots are not real.",
        "solution_hi": "D = (-2)² - 4(1)(5) = -16। क्योंकि D < 0 है, इसलिए वास्तविक मूल नहीं हैं।",
        "marking": "1 mark for D = -16; 1 mark for the correct conclusion.",
    },
    {
        "n": 14, "section": "short", "marks": 2,
        "en": "Solve 3x² = 5x + 2 and verify both roots in the original equation.",
        "hi": "3x² = 5x + 2 को हल कीजिए और दोनों मूलों की मूल समीकरण में जाँच कीजिए।",
        "answer": "x = 2 or x = -1/3",
        "solution_en": "3x² - 5x - 2 = (3x + 1)(x - 2) = 0. Thus x = -1/3 or 2; substitution satisfies the original equation in both cases.",
        "solution_hi": "3x² - 5x - 2 = (3x + 1)(x - 2) = 0। अतः x = -1/3 अथवा 2; दोनों मान मूल समीकरण को संतुष्ट करते हैं।",
        "marking": "1 mark for both roots; 1 mark for a valid verification.",
    },
    {
        "n": 15, "section": "application", "marks": 3,
        "en": "A rectangular study board has area 54 dm². Its length is 3 dm more than its width. Find its dimensions.",
        "hi": "एक आयताकार अध्ययन-पट्ट का क्षेत्रफल 54 डीएम² है। उसकी लंबाई चौड़ाई से 3 डीएम अधिक है। उसके आयाम ज्ञात कीजिए।",
        "answer": "Width 6 dm; length 9 dm",
        "solution_en": "Let width = x, so length = x + 3. Then x(x + 3) = 54, or (x + 9)(x - 6) = 0. Reject x = -9. Dimensions are 6 dm by 9 dm.",
        "solution_hi": "चौड़ाई x और लंबाई x + 3 मानें। x(x + 3) = 54 से (x + 9)(x - 6) = 0। ऋणात्मक मान अस्वीकार करने पर आयाम 6 डीएम और 9 डीएम हैं।",
        "marking": "1 mark for model; 1 mark for solving; 1 mark for valid dimensions with rejection of the negative root.",
    },
    {
        "n": 16, "section": "application", "marks": 3,
        "en": "A ball's height above the ground is h(t) = -5t² + 20t + 1 metres, t seconds after launch. When does it first return to the ground? Give the answer to 2 decimal places.",
        "hi": "प्रक्षेपण के t सेकंड बाद गेंद की जमीन से ऊँचाई h(t) = -5t² + 20t + 1 मीटर है। गेंद पहली बार जमीन पर कब लौटेगी? उत्तर 2 दशमलव स्थान तक दीजिए।",
        "answer": "t = 2 + √105/5 ≈ 4.05 s",
        "solution_en": "Set h = 0: 5t² - 20t - 1 = 0. Thus t = [20 ± √420]/10 = 2 ± √105/5. Time cannot be negative, so t ≈ 4.05 s.",
        "solution_hi": "h = 0 रखने पर 5t² - 20t - 1 = 0। अतः t = 2 ± √105/5। ऋणात्मक समय अस्वीकार करने पर t ≈ 4.05 सेकंड।",
        "marking": "1 mark for equation; 1 mark for exact roots; 1 mark for 4.05 s with negative root rejected.",
    },
    {
        "n": 17, "section": "application", "marks": 3,
        "en": "The product of two consecutive positive integers is 156. Find the integers and show why the negative solution is inadmissible.",
        "hi": "दो क्रमागत धन पूर्णांकों का गुणनफल 156 है। पूर्णांक ज्ञात कीजिए और बताइए कि ऋणात्मक हल स्वीकार्य क्यों नहीं है।",
        "answer": "12 and 13",
        "solution_en": "Let the integers be n and n + 1. Then n² + n - 156 = (n + 13)(n - 12) = 0. Since the integers are positive, n = 12; the integers are 12 and 13.",
        "solution_hi": "पूर्णांक n और n + 1 मानें। n² + n - 156 = (n + 13)(n - 12) = 0। धन पूर्णांक होने से n = 12; पूर्णांक 12 और 13 हैं।",
        "marking": "1 mark for model; 1 mark for solving; 1 mark for answer and domain reasoning.",
    },
    {
        "n": 18, "section": "application", "marks": 3,
        "en": "Find the exact points where y = x² - 4x + 3 intersects y = 2x - 1.",
        "hi": "y = x² - 4x + 3 और y = 2x - 1 के प्रतिच्छेद बिंदु सटीक रूप में ज्ञात कीजिए।",
        "answer": "(3 + √5, 5 + 2√5) and (3 - √5, 5 - 2√5)",
        "solution_en": "Equate: x² - 4x + 3 = 2x - 1, so x² - 6x + 4 = 0. Hence x = 3 ± √5. Using y = 2x - 1 gives y = 5 ± 2√5 with matching signs.",
        "solution_hi": "दोनों y बराबर रखने पर x² - 6x + 4 = 0, इसलिए x = 3 ± √5। y = 2x - 1 में रखने पर y = 5 ± 2√5 मिलता है।",
        "marking": "1 mark for intersection equation; 1 mark for x-values; 1 mark for both ordered pairs.",
    },
    {
        "n": 19, "section": "extended", "marks": 4,
        "en": "A school makes a rectangular garden against a straight wall, so fencing is needed on only three sides. There are 40 m of fencing. Model the area, then determine the dimensions that give the maximum area and justify the maximum.",
        "hi": "एक विद्यालय सीधी दीवार के सहारे आयताकार बगीचा बनाता है, इसलिए केवल तीन भुजाओं पर बाड़ चाहिए। कुल 40 मीटर बाड़ उपलब्ध है। क्षेत्रफल का मॉडल बनाइए, अधिकतम क्षेत्रफल देने वाले आयाम ज्ञात कीजिए और अधिकतम होने का औचित्य दीजिए।",
        "answer": "Width 10 m, length 20 m; maximum area 200 m²",
        "solution_en": "Let each perpendicular side be x, so the parallel side is 40 - 2x. A(x) = x(40 - 2x) = -2(x - 10)² + 200. The squared term is at least 0, so the maximum is 200 m² at x = 10; dimensions are 10 m by 20 m.",
        "solution_hi": "दीवार पर लंबवत प्रत्येक भुजा x हो, तो तीसरी भुजा 40 - 2x होगी। A(x) = x(40 - 2x) = -2(x - 10)² + 200। वर्ग पद न्यूनतम 0 होने पर अधिकतम क्षेत्रफल 200 मी² तथा आयाम 10 मी और 20 मी हैं।",
        "marking": "1 mark for constraint/model; 1 mark for completed-square or vertex method; 1 mark for dimensions; 1 mark for maximum-area justification.",
    },
    {
        "n": 20, "section": "extended", "marks": 4,
        "en": "A student solves x² + 6x + 10 = 0 as follows: x² + 6x = -10; (x + 3)² = 1; x = -2 or -4. Identify the error, give the correct conclusion over the real numbers, and connect your conclusion to the graph.",
        "hi": "एक विद्यार्थी x² + 6x + 10 = 0 को इस प्रकार हल करता है: x² + 6x = -10; (x + 3)² = 1; x = -2 या -4। त्रुटि पहचानिए, वास्तविक संख्याओं में सही निष्कर्ष दीजिए और उसे ग्राफ से जोड़कर समझाइए।",
        "answer": "The completed square is (x + 3)² = -1; no real roots; vertex (-3, 1) lies above the x-axis",
        "solution_en": "Adding 9 to both sides gives (x + 3)² = -1, not 1. No real number has square -1, so there are no real roots. Equivalently y = (x + 3)² + 1 has vertex (-3, 1) and never meets the x-axis.",
        "solution_hi": "दोनों पक्षों में 9 जोड़ने पर (x + 3)² = -1 मिलता है, 1 नहीं। किसी वास्तविक संख्या का वर्ग -1 नहीं होता, इसलिए वास्तविक मूल नहीं हैं। y = (x + 3)² + 1 का शीर्ष (-3, 1) है और ग्राफ x-अक्ष को नहीं काटता।",
        "marking": "1 mark for locating the sign error; 1 mark for corrected square; 1 mark for no-real-root conclusion; 1 mark for the graph/vertex explanation.",
    },
]


SECTION_META = {
    "mcq": ("SECTION A · CONCEPT CHECK", "खंड A · अवधारणा जाँच", "8 questions × 1 mark = 8"),
    "short": ("SECTION B · METHOD & REASONING", "खंड B · विधि एवं तर्क", "6 questions × 2 marks = 12"),
    "application": ("SECTION C · APPLICATION", "खंड C · अनुप्रयोग", "4 questions × 3 marks = 12"),
    "extended": ("SECTION D · TRANSFER & CRITIQUE", "खंड D · स्थानांतरण एवं समीक्षा", "2 questions × 4 marks = 8"),
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    total = sum(widths_dxa)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_table_no_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "nil")


def set_run_font(run, *, size: float = 10.5, bold: bool = False, color: str = INK, hindi: bool = False, italic: bool = False) -> None:
    font_name = DEVANAGARI_FONT if hindi else LATIN_FONT
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), LATIN_FONT)
    r_fonts.set(qn("w:hAnsi"), LATIN_FONT)
    r_fonts.set(qn("w:cs"), DEVANAGARI_FONT)
    r_fonts.set(qn("w:eastAsia"), DEVANAGARI_FONT)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "hi-IN" if hindi else "en-IN")
    lang.set(qn("w:bidi"), "hi-IN")


def set_paragraph_spacing(paragraph, *, before: float = 0, after: float = 5, line: float = 1.15, keep: bool = False) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line
    fmt.keep_together = keep


def add_text(paragraph, text: str, *, hindi: bool = False, size: float = 10.5, bold: bool = False, color: str = INK, italic: bool = False):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, hindi=hindi, italic=italic)
    return run


def add_para(doc, text: str = "", *, hindi: bool = False, size: float = 10.5, bold: bool = False, color: str = INK, before: float = 0, after: float = 5, line: float = 1.15, align=None, keep: bool = False, italic: bool = False):
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    set_paragraph_spacing(paragraph, before=before, after=after, line=line, keep=keep)
    if text:
        add_text(paragraph, text, hindi=hindi, size=size, bold=bold, color=color, italic=italic)
    return paragraph


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def set_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_text(paragraph, "PAGE ", size=8.5, bold=True, color=MUTED)
    run = paragraph.add_run()
    set_run_font(run, size=8.5, bold=True, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def style_document(doc: Document, title: str, subject: str) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), DEVANAGARI_FONT)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in (
        ("Title", 26, INK, 0, 6),
        ("Subtitle", 12, INK_SOFT, 0, 12),
        ("Heading 1", 16, INK, 14, 8),
        ("Heading 2", 12, BLUE, 10, 5),
        ("Heading 3", 11, INK_SOFT, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = LATIN_FONT
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name != "Subtitle"
        style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
        style._element.rPr.rFonts.set(qn("w:cs"), DEVANAGARI_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    props = doc.core_properties
    props.title = title
    props.subject = subject
    props.author = "Teacher Prompt Studio"
    props.last_modified_by = "Teacher Prompt Studio"
    props.keywords = "TPS-IY-79; en-IN; hi-IN; Class 10; Quadratic Equations"
    props.comments = "Curated and render-verified flagship sample."

    # Keep the running header empty. LibreOffice can mirror complex-script run
    # language properties unpredictably in even-page headers, so the visible
    # provenance line lives in the first-page title block and metadata instead.
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = ""

    footer = section.footer
    footer_table = footer.add_table(rows=1, cols=2, width=Cm(17))
    set_table_geometry(footer_table, [7000, 2360], indent_dxa=0)
    footer_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    left = footer_table.cell(0, 0).paragraphs[0]
    set_paragraph_spacing(left, after=0)
    add_text(left, "Classroom practice sample · Not an official board paper · TPS-IY-79", size=8, color=MUTED)
    right = footer_table.cell(0, 1).paragraphs[0]
    set_page_number(right)


def add_title_block(doc: Document, *, teacher: bool) -> None:
    brand = add_para(doc, "TEACHER PROMPT STUDIO  /  FLAGSHIP FIELD SAMPLE", size=8, bold=True, color=MUTED, after=10, keep=True)
    brand.paragraph_format.keep_with_next = True
    kicker = "TEACHER-ONLY ASSESSMENT PACK" if teacher else "CLASSROOM ASSESSMENT · STUDENT PAPER"
    paragraph = add_para(doc, kicker, size=8.5, bold=True, color=BLUE, after=4)
    paragraph.paragraph_format.keep_with_next = True
    title = add_para(doc, "QUADRATIC EQUATIONS", size=25, bold=True, color=INK, after=1, keep=True)
    add_text(title, "  /  द्विघात समीकरण", hindi=True, size=17, bold=True, color=INK_SOFT)
    subtitle = "Worked key, marking logic and blueprint" if teacher else "Bilingual assessment · English + Hindi"
    subtitle_hi = "हल, अंकन-तर्क और रूपरेखा" if teacher else "द्विभाषी आकलन · अंग्रेज़ी + हिन्दी"
    p = add_para(doc, subtitle, size=11.5, color=INK_SOFT, after=2, keep=True)
    add_text(p, f"  ·  {subtitle_hi}", hindi=True, size=10.5, color=INK_SOFT)

    table = doc.add_table(rows=2, cols=4)
    set_table_geometry(table, [1500, 3180, 1500, 3180])
    labels = [("CLASS", "10"), ("TIME", "60 min"), ("MARKS", "40"), ("VERSION", "EN + HI")]
    for index, (label, value) in enumerate(labels):
        row = index // 2
        col = (index % 2) * 2
        set_cell_shading(table.cell(row, col), INK)
        p_label = table.cell(row, col).paragraphs[0]
        set_paragraph_spacing(p_label, after=0)
        add_text(p_label, label, size=8, bold=True, color=LIME)
        p_value = table.cell(row, col + 1).paragraphs[0]
        set_paragraph_spacing(p_value, after=0)
        add_text(p_value, value, size=10.5, bold=True, color=INK)
    add_para(doc, "CBSE / NCERT chapter context · teacher-created practice sample; no official paper pattern is implied.", size=8.5, color=MUTED, after=10)


def add_student_identity(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [3800, 2780, 2780])
    for cell, text in zip(table.rows[0].cells, ["Name / नाम", "Roll no. / अनुक्रमांक", "Date / दिनांक"]):
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, after=13)
        add_text(p, text, hindi="नाम" in text or "अनुक्रमांक" in text or "दिनांक" in text, size=9, bold=True, color=INK_SOFT)
        p_pr = p._p.get_or_add_pPr()
        border = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "5")
        bottom.set(qn("w:color"), LINE)
        border.append(bottom)
        p_pr.append(border)
    add_para(doc, after=3)


def add_section_banner(doc: Document, section_id: str, *, page_break_before: bool = False):
    en, hi, marks = SECTION_META[section_id]
    p = add_para(doc, before=10, after=3, keep=True)
    p.paragraph_format.page_break_before = page_break_before
    set_paragraph_shading(p, INK)
    add_text(p, f"  {en}", size=10.5, bold=True, color=WHITE)
    add_text(p, f"   {hi}", hindi=True, size=10, bold=True, color=LIME)
    q = add_para(doc, marks, size=8.5, bold=True, color=BLUE, after=7, keep=True)
    q.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return p


def add_continuation_banner(doc: Document, section_id: str, *, page_break_before: bool = False):
    en, hi, _ = SECTION_META[section_id]
    p = add_para(doc, before=0, after=8, keep=True)
    p.paragraph_format.page_break_before = page_break_before
    set_paragraph_shading(p, INK)
    add_text(p, f"  {en} · CONTINUED", size=10, bold=True, color=WHITE)
    add_text(p, f"   {hi} · जारी", hindi=True, size=9.6, bold=True, color=LIME)
    return p


def add_question(doc, item: dict, *, teacher: bool, table_width: int = 9360, table_indent: int = 120) -> None:
    p = add_para(doc, before=5, after=2, keep=True)
    set_keep_with_next(p)
    add_text(p, f"Q{item['n']}. ", size=10.5, bold=True, color=BLUE)
    add_text(p, item["en"], size=10.5, bold=True, color=INK)
    add_text(p, f"  [{item['marks']}]", size=9, bold=True, color=MUTED)
    h = add_para(doc, item["hi"], hindi=True, size=10.2, color=INK_SOFT, after=4, keep=True)
    h.paragraph_format.left_indent = Cm(0.45)

    if item.get("options"):
        table = doc.add_table(rows=2, cols=2)
        set_table_geometry(table, [table_width // 2, table_width // 2], indent_dxa=table_indent)
        for index, option in enumerate(item["options"]):
            cell = table.cell(index // 2, index % 2)
            p_opt = cell.paragraphs[0]
            set_paragraph_spacing(p_opt, after=0)
            add_text(p_opt, option, size=9.8, color=INK)
        if teacher:
            for index, option in enumerate(item["options"]):
                if option.startswith(f"{item['answer']}."):
                    set_cell_shading(table.cell(index // 2, index % 2), MINT)
    if teacher:
        callout = doc.add_table(rows=1, cols=1)
        set_table_geometry(callout, [table_width], indent_dxa=table_indent)
        cant_split = OxmlElement("w:cantSplit")
        callout.rows[0]._tr.get_or_add_trPr().append(cant_split)
        set_cell_shading(callout.cell(0, 0), PALE_BLUE)
        cell = callout.cell(0, 0)
        p_answer = cell.paragraphs[0]
        set_paragraph_spacing(p_answer, after=3)
        add_text(p_answer, f"ANSWER  {item['answer']}", size=9, bold=True, color=BLUE)
        p_solution = cell.add_paragraph()
        set_paragraph_spacing(p_solution, after=2)
        add_text(p_solution, item["solution_en"], size=9.5, color=INK)
        p_hi = cell.add_paragraph()
        set_paragraph_spacing(p_hi, after=2)
        add_text(p_hi, item["solution_hi"], hindi=True, size=9.2, color=INK_SOFT)
        p_mark = cell.add_paragraph()
        set_paragraph_spacing(p_mark, after=0)
        add_text(p_mark, f"MARKING  {item['marking']}", size=8.5, bold=True, color=MUTED)
    elif item["section"] != "mcq":
        heights = {"short": 1.05, "application": 2.05, "extended": 3.25}
        response = doc.add_table(rows=1, cols=1)
        response.style = "Table Grid"
        set_table_geometry(response, [table_width], indent_dxa=table_indent)
        row = response.rows[0]
        row.height = Cm(heights[item["section"]])
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        cell = response.cell(0, 0)
        set_cell_shading(cell, "FBFCFC")
        p_response = cell.paragraphs[0]
        set_paragraph_spacing(p_response, after=0)
        add_text(p_response, "RESPONSE  /  उत्तर", hindi=True, size=7.8, bold=True, color=MUTED)
    add_para(doc, after=2)


def add_instructions(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_cell_shading(table.cell(0, 0), PALE_GOLD)
    cell = table.cell(0, 0)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=4)
    add_text(p, "BEFORE YOU BEGIN  /  आरम्भ करने से पहले", hindi=True, size=9.5, bold=True, color=INK)
    instructions = [
        ("Attempt all 20 questions. Show complete working where the question asks for a method or justification.", "सभी 20 प्रश्न हल कीजिए। जहाँ विधि या औचित्य माँगा गया है, पूरा कार्य दिखाइए।"),
        ("Choose exactly one option in Section A. Calculators are not required.", "खंड A में केवल एक विकल्प चुनिए। कैलकुलेटर की आवश्यकता नहीं है।"),
        ("Keep exact values unless a question asks for decimal rounding.", "जब तक दशमलव सन्निकटन न माँगा जाए, सटीक मान लिखिए।"),
    ]
    for en, hi in instructions:
        para = cell.add_paragraph()
        set_paragraph_spacing(para, after=3)
        add_text(para, "• ", size=9.2, bold=True, color=BLUE)
        add_text(para, en, size=9.2, color=INK)
        add_text(para, f"  {hi}", hindi=True, size=9, color=INK_SOFT)
    add_para(doc, after=5)


def build_student_doc() -> Document:
    doc = Document()
    style_document(doc, "Class 10 Quadratic Equations - Bilingual Student Paper", "Mathematics classroom assessment")
    add_title_block(doc, teacher=False)
    add_student_identity(doc)
    add_instructions(doc)
    current = None
    for item in ITEMS:
        if item["section"] != current:
            current = item["section"]
            add_section_banner(doc, current)
        if item["n"] in {4, 10, 16, 20}:
            doc.add_page_break()
            add_continuation_banner(doc, item["section"])
        add_question(doc, item, teacher=False)
    p = add_para(doc, "END OF PAPER  /  प्रश्नपत्र समाप्त", hindi=True, size=9.5, bold=True, color=BLUE, before=12, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_paragraph_shading(p, PALE_BLUE)
    return doc


def add_blueprint(doc: Document) -> None:
    add_section_banner(doc, "mcq")
    p = add_para(doc, "Assessment blueprint · आकलन रूपरेखा", hindi=True, size=14, bold=True, color=INK, before=4, after=6)
    set_keep_with_next(p)
    table = doc.add_table(rows=1, cols=5)
    set_table_geometry(table, [2900, 1600, 1300, 1500, 2060])
    headers = ["Demand", "Items", "Marks each", "Total", "Evidence"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, INK)
        p_head = cell.paragraphs[0]
        set_paragraph_spacing(p_head, after=0)
        add_text(p_head, text, size=8.5, bold=True, color=LIME)
    rows = [
        ("Concept check", "Q1-Q8", "1", "8", "Breadth + discrimination"),
        ("Method & reasoning", "Q9-Q14", "2", "12", "Procedure + explanation"),
        ("Application", "Q15-Q18", "3", "12", "Modelling + exact analysis"),
        ("Transfer & critique", "Q19-Q20", "4", "8", "Justification + evaluation"),
        ("TOTAL", "20 items", "-", "40", "60 minutes"),
    ]
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for cell, text in zip(cells, values):
            p_cell = cell.paragraphs[0]
            set_paragraph_spacing(p_cell, after=0)
            add_text(p_cell, text, size=8.5, bold=row_index == len(rows) - 1, color=INK)
        if row_index == len(rows) - 1:
            for cell in cells:
                set_cell_shading(cell, MINT)
    set_table_geometry(table, [2900, 1600, 1300, 1500, 2060])
    add_para(doc, "Reasoning marks: 32/40 (80%). The paper is a curated classroom sample, not a claim of official board-pattern alignment.", size=8.7, color=MUTED, after=8)


def add_teacher_protocol(doc: Document) -> None:
    p = add_para(doc, "RELEASE EVIDENCE", size=11, bold=True, color=BLUE, before=8, after=5)
    set_keep_with_next(p)
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [2300, 7060])
    evidence = [
        ("FILES", "Student PDF + editable student DOCX + separate teacher PDF"),
        ("CONTENT", "20/20 questions; 8 + 6 + 4 + 2 distribution; 40/40 marks"),
        ("AUDIENCE", "No answers, blueprint, assumptions or marking guidance in either student file"),
        ("LANGUAGE", "English + Hindi text preserved through DOCX and PDF render"),
        ("STATUS", "Curated sample; teacher should still review local wording before classroom use"),
    ]
    for index, (label, value) in enumerate(evidence):
        cells = table.rows[0].cells if index == 0 else table.add_row().cells
        set_cell_shading(cells[0], INK)
        p_label = cells[0].paragraphs[0]
        set_paragraph_spacing(p_label, after=0)
        add_text(p_label, label, size=8.3, bold=True, color=LIME)
        p_value = cells[1].paragraphs[0]
        set_paragraph_spacing(p_value, after=0)
        add_text(p_value, value, size=8.7, color=INK)
    set_table_geometry(table, [2300, 7060])
    add_para(doc, after=4)


def build_teacher_doc() -> Document:
    doc = Document()
    style_document(doc, "Class 10 Quadratic Equations - Bilingual Teacher Pack", "Blueprint, worked key and marking guidance")
    add_title_block(doc, teacher=True)
    warning = add_para(doc, "TEACHER ONLY · DO NOT DISTRIBUTE WITH THE STUDENT PAPER  /  केवल शिक्षक हेतु", hindi=True, size=9.5, bold=True, color=INK, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_paragraph_shading(warning, PALE_GOLD)
    add_blueprint(doc)
    add_teacher_protocol(doc)
    p = add_para(doc, "WORKED ANSWER & MARKING KEY  /  हल एवं अंकन-कुंजी", hindi=True, size=16, bold=True, color=INK, before=14, after=8)
    p.paragraph_format.page_break_before = True
    current = None
    for item in ITEMS:
        if item["section"] != current:
            new_page = current is not None
            current = item["section"]
            add_section_banner(doc, current, page_break_before=new_page)
        elif item["n"] in {4, 7, 13, 18}:
            add_continuation_banner(doc, item["section"], page_break_before=True)
        block = doc.add_table(rows=1, cols=1)
        set_table_geometry(block, [9360])
        set_table_no_borders(block)
        row = block.rows[0]
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        cell = block.cell(0, 0)
        set_cell_margins(cell, top=0, start=0, bottom=0, end=0)
        empty = cell.paragraphs[0]._element
        empty.getparent().remove(empty)
        add_question(cell, item, teacher=True, table_width=9120, table_indent=0)
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(1)
        spacer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        spacer.paragraph_format.line_spacing = Pt(1)
    return doc


def validate_content() -> None:
    assert [item["n"] for item in ITEMS] == list(range(1, 21))
    counts = Counter(item["section"] for item in ITEMS)
    assert counts == Counter({"mcq": 8, "short": 6, "application": 4, "extended": 2})
    assert sum(item["marks"] for item in ITEMS) == 40
    assert all(len(item.get("options", [])) == 4 for item in ITEMS if item["section"] == "mcq")
    assert all(item.get("answer") and item.get("solution_en") and item.get("solution_hi") and item.get("marking") for item in ITEMS)
    corpus = "\n".join(str(value) for item in ITEMS for value in item.values())
    for forbidden in ("TBD", "TODO", "FIXME", "Lorem ipsum", "{{", "[INSERT", "..."):
        assert forbidden not in corpus
    assert "�" not in corpus
    assert sum(1 for char in corpus if "\u0900" <= char <= "\u097f") > 500


def save_doc(doc: Document, path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)
    assert path.exists() and path.stat().st_size > 20_000


def main() -> None:
    WORK_DIR.mkdir(parents=True, exist_ok=True)
    PUBLIC_DIR.mkdir(parents=True, exist_ok=True)
    validate_content()
    save_doc(build_student_doc(), STUDENT_DOCX)
    save_doc(build_teacher_doc(), TEACHER_SOURCE)
    print(f"student_docx={STUDENT_DOCX}")
    print(f"teacher_source={TEACHER_SOURCE}")
    print("items=20 marks=40 mix=8/6/4/2")


if __name__ == "__main__":
    main()
